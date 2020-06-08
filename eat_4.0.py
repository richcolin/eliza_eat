from  docx import  Document
from  docx.shared import  Pt
from  docx.oxml.ns import  qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import os
import glob
import random
print("饭费原因文件名为reason.txt.\n具体问题文件名为eat.txt")
single_price=int(input('请输入单价：'))
keep_name=int(input('固定人名的个数：'))
reason_num=int(input('请输入理由个数：'))
path = os.getcwd()
local_txt = glob.glob(os.path.join(path, 'eat.txt'))
print(local_txt)
local_txt_reason=glob.glob(os.path.join(path, 'reason.txt'))
print(local_txt_reason)
local_txt=str(local_txt[0])
local_txt_reason=str(local_txt_reason[0])
f_reason=open(local_txt_reason, mode='r', encoding='utf8')

f = open(local_txt, mode='r', encoding='utf8')

reason_list=[]
eat_dict={}
name_dict={}
contact_dict={}
for line in f:
    line=line.strip('\n')
    each_list=line.split(' ')
    if each_list[0]=='姓名':
        eat_dict[each_list[0]]=each_list[1::]
    else:

        if each_list[0] in eat_dict.keys():

            eat_dict[each_list[0]] =int(each_list[1::][0])+eat_dict[each_list[0]]
        else:
            eat_dict[each_list[0]] = int(each_list[1::][0])
for line in f_reason:
    line=line.strip('\n')
    reason_list.append(line)
print(reason_list)
for key,value in eat_dict.items():

    if key!='姓名':

        contact_dict[key] =value
    else:
        name_dict[key]=value

produce_dict={}
total_price=0
print('contact',contact_dict)
for key,value in contact_dict.items():

    people_num=int(value/single_price)

    if people_num<=keep_name:

        produce_dict[key] = random.sample(set(name_dict['姓名'][0:keep_name]), people_num)
    else:

        produce_dict[key]=random.sample(set(name_dict['姓名']), people_num)
    total_price=total_price+value
def zhengze(jianzhi):
    # 利用正则表达式
    import re
    data = jianzhi
    patt = '(\d+).(\d+).(\d+)'
    # 交换排序
    for i in range(len(data) - 1):
        for x in range(i + 1, len(data)):
            j = 1
            while j < 4:
                lower = re.match(patt, data[i]).group(j)
                upper = re.match(patt, data[x]).group(j)
                # print lower,upper
                if int(lower) < int(upper):
                    j = 4
                elif int(lower) == int(upper):
                    j += 1
                else:
                    data[i], data[x] = data[x], data[i]
                    j = 4

    return data
keys=list(produce_dict.keys())
numerd_keys=zhengze(keys)
brand_new_dict={}
for datess in keys:
    brand_new_dict[datess]=produce_dict[datess]
print(brand_new_dict,'brand')
print(numerd_keys,'nameerd')
print(produce_dict)


ordered_key=brand_new_dict
f.close()

document = Document()

paragraph = document.add_paragraph()
paragraph.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
run = paragraph.add_run(u'用餐情况说明')
run.bold=True
run.font.name=u'宋体'
r = run._element
r.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
run.font.size=Pt(16)

paragraph = document.add_paragraph()
run.font.name=u'宋体'
r = run._element
r.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
run.font.size=Pt(14)
for reason_single in random.sample(reason_list, reason_num):

    paragraph = document.add_paragraph()
    run = paragraph.add_run(reason_single)
    run.font.name=u'宋体'
    r = run._element
    r.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
    run.font.size=Pt(14)




paragraph = document.add_paragraph()
run = paragraph.add_run(u'名单如下：')
run.font.name=u'宋体'
r = run._element
r.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
run.font.size=Pt(14)
for order_key in ordered_key:

    paragraph = document.add_paragraph()
    run = paragraph.add_run(order_key)
    run.font.name=u'宋体'
    r = run._element
    r.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
    run.font.size=Pt(14)
    for v in produce_dict[order_key]:

        run = paragraph.add_run(" %s"%v)
        run.font.name = u'宋体'
        r = run._element
        r.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
        run.font.size = Pt(14)

paragraph = document.add_paragraph()
paragraph.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
run = paragraph.add_run('共计：%d'%total_price)
run.font.name=u'宋体'
r = run._element
r.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
run.font.size=Pt(14)

#增加分页
document.add_page_break()
from datetime import datetime
now = datetime.now()
date_now=now.date()
#保存文件
document.save('%s餐费统计.docx'%date_now)
